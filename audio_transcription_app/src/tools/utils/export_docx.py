import os
from typing import List
from docx import Document

def export_docx(transcription: str, output_path: str, title: str = "Transcription"):
    doc = Document()
    doc.add_heading(title, 0)
    for line in transcription.splitlines():
        doc.add_paragraph(line)
    doc.save(output_path)
